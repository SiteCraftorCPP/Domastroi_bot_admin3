import asyncio
import os
import logging
from datetime import datetime
from aiogram import Bot, Dispatcher, types
from aiogram.contrib.fsm_storage.memory import MemoryStorage
from aiogram.dispatcher import FSMContext
from aiogram.dispatcher.filters.state import State, StatesGroup
from aiogram.utils.exceptions import MessageNotModified
from aiogram.dispatcher.filters import Text
from aiogram.utils.executor import start_polling
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from io import BytesIO
import requests
import asyncpg
import json

# Загрузка переменных окружения
load_dotenv()

# Настройка логирования
logging.basicConfig(level=logging.INFO)

DB_USER = os.getenv('DB_USER')
DB_PASSWORD = os.getenv('DB_PASSWORD')
DB_NAME = os.getenv('DB_NAME')
DB_HOST = os.getenv('DB_HOST')
DB_PORT = os.getenv('DB_PORT', '5432')

if not DB_USER or not DB_PASSWORD or not DB_NAME or not DB_HOST:
    raise ValueError("Не все переменные окружения загружены корректно")

# Подключение к базе данных
async def create_db_pool():
    return await asyncpg.create_pool(
        user=DB_USER,
        password=DB_PASSWORD,
        database=DB_NAME,
        host=DB_HOST,
        port=DB_PORT
    )

db_pool = None

def load_questions():
    with open('questions.json', 'r', encoding='utf-8') as file:
        data = json.load(file)
    return data["questions"]

# Функция для получения ID папки
def get_folder_id():
    current_path = os.path.dirname(os.path.realpath(__file__))  # Получаем текущий путь к файлу
    folder_name = os.path.basename(current_path)  # Извлекаем имя последней папки в пути
    try:
        # Пробуем преобразовать строковый ID в целочисленный
        folder_id_int = int(folder_name)
        return folder_id_int
    except ValueError:
        # В случае ошибки преобразования возвращаем None или дефолтное значение
        return None

# Получение токена бота из базы данных, используя ID папки
async def get_bot_token():
    folder_id = get_folder_id()  # Получаем ID из названия папки
    try:
        # Пробуем преобразовать строковый ID в целочисленный
        folder_id_int = int(folder_id)
    except ValueError:
        # В случае ошибки преобразования выводим сообщение и завершаем функцию
        print("ID папки должен быть числом.")
        return None

    async with db_pool.acquire() as connection:
        result = await connection.fetchval(
            """
            SELECT bot_api FROM users WHERE id_telegram = $1
            """,
            folder_id_int  # Используем преобразованный целочисленный ID
        )
        return result


# Проверка срока действия подписки
async def check_subscription(user_id):
    async with db_pool.acquire() as connection:
        # Получаем ID из названия папки
        folder_id = get_folder_id()
        
        # Проверяем статус подписки
        subscription_status = await connection.fetchrow(
            """
            SELECT pay, date_stop FROM users WHERE id_telegram = $1
            """,
            folder_id
        )
        
        if subscription_status is None:
            return {'status': False, 'message': "Настройки бота не настроены корректно или аккаунт отсутствует."}
        
        if subscription_status['pay'] == 1 and subscription_status['date_stop'] > datetime.now():
            return {'status': True}
        else:
            return {'status': False, 'message': "Владелец бота не продлил свою подписку, функционал временно ограничен.\nБот разработан @WSCHUDO"}


# Определение состояний FSM
class Form(StatesGroup):
    example_state = State()

class Form(StatesGroup):
    waiting_for_phone = State()
    waiting_for_phone_confirmation = State()

class Questionnaire(StatesGroup):
    asking = State()  # Состояние для всех вопросов
    custom_answer = State()  # Состояние для пользовательского ответа


# Основная функция
async def main():
    global db_pool
    db_pool = await create_db_pool()
    
    API_TOKEN = await get_bot_token()
    if not API_TOKEN:
        raise ValueError("Токен бота не найден в базе данных")

    # Создание экземпляра бота и диспетчера
    global bot
    bot = Bot(token=API_TOKEN)
    storage = MemoryStorage()
    dp = Dispatcher(bot, storage=storage)

    # Проверка на нпродолжение 

    # Функция для проверки последнего шага пользователя
    async def check_last_step(user_id):
        pool = await create_db_pool()
        async with pool.acquire() as connection:
            last_step = await connection.fetchval(
                """
                SELECT status FROM users_designer WHERE id_telegram = $1
                """,
                user_id
            )
        await pool.close()
        return last_step


    # Меню по команде menu

    # # Функция для проверки статуса подписки
    # async def check_subscription_status(user_id):
    #     pool = await create_db_pool()
    #     async with pool.acquire() as connection:
    #         row = await connection.fetchrow(
    #             """
    #             SELECT pay, date_stop FROM users WHERE id_telegram = $1
    #             """,
    #             user_id
    #         )
    #     await pool.close()

    #     if row:
    #         return {'pay': row['pay'], 'date_stop': row['date_stop']}
    #     else:
    #         # Возвращаем значение по умолчанию, если запись не найдена
    #         return {'pay': 0, 'date_stop': datetime.min}

    @dp.message_handler(commands=['menu'])
    @dp.message_handler(lambda message: message.text.lower() == 'меню')
    async def show_menu(message: types.Message):

        subscription_status = await check_subscription(message.from_user.id)

        if subscription_status['status']:

            last_step = await check_last_step(message.from_user.id)
            # Создание клавиатуры для взаимодействия с ботом
            keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
            if last_step != 0:
                # buttons = [
                #     ["Продолжить", "Портфолио"],
                #     ["Хелпер", "Контакты"]
                # ]
                buttons = [
                    ["Продолжить", "Хелпер"]
                ]
            else:
                # buttons = [
                #     ["Заполнить ТЗ", "Портфолио"],
                #     ["Хелпер", "Контакты"]
                # ]
                buttons = [
                    ["Заполнить ТЗ", "Хелпер"]
                ]
            for row in buttons:
                keyboard.row(*row)

            await message.answer("Выберите необходимый пункт меню", reply_markup=keyboard)

        else:
            await message.answer(subscription_status['message'])    

    # END Меню по команде menu

    # Пример обработчика команды /start
    @dp.message_handler(commands='start')
    async def cmd_start(message: types.Message):
        folder_id = get_folder_id()  # Получаем ID из названия папки
        if folder_id is None:
            await message.answer("Ошибка конфигурации бота.")
            return
        async with db_pool.acquire() as connection:
            root_user = await connection.fetchrow(
                """
                SELECT id_telegram FROM users WHERE id_telegram = $1
                """,
                folder_id
            )
            await connection.execute(
                """
                INSERT INTO users_designer (id_telegram, tg_login, tg_firstname, tg_lastname, status, phone, last_step, subscribe, root)
                VALUES ($1, $2, $3, $4, 0, NULL, 0, 0, $5)
                ON CONFLICT (id_telegram) DO NOTHING
                """,
                message.from_user.id, message.from_user.username, message.from_user.first_name, message.from_user.last_name, root_user['id_telegram'] if root_user else folder_id
            )
            # Проверяем статус подписки для ID из названия папки
            subscription_status = await connection.fetchrow(
                """
                SELECT pay, date_stop FROM users WHERE id_telegram = $1
                """,
                folder_id
            )
            
            if subscription_status is None:
                await message.answer("Настройки бота не настроены корректно или аккаунт отсутствует.")
                return

            if subscription_status['pay'] == 1 and subscription_status['date_stop'] > datetime.now():
                # Создание клавиатуры для взаимодействия с ботом
                last_step = await check_last_step(message.from_user.id)
                # Создание клавиатуры для взаимодействия с ботом
                keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
                if last_step != 0:
                    # buttons = [
                    #     ["Продолжить", "Портфолио"],
                    #     ["Хелпер", "Контакты"]
                    # ]
                    buttons = [
                        ["Продолжить", "Хелпер"]
                    ]
                else:
                    # buttons = [
                    #     ["Заполнить ТЗ", "Портфолио"],
                    #     ["Хелпер", "Контакты"]
                    # ]
                    buttons = [
                        ["Заполнить ТЗ", "Хелпер"]
                    ]
                for row in buttons:
                    keyboard.row(*row)

                # await message.answer("Добро пожаловать в нашего помощника по заполнению технических заданий!\n\n📝 Заполнить ТЗ — Нажмите /technical_specification, чтобы начать или продолжить\n\n🎨 Портфолио — Посмотрите наши проекты и работы /portfolio\n\n🤖 Хелпер — Получите помощь и советы по использованию бота /help\n\n📞 Контакты — Свяжитесь с нами для получения дополнительной информации /contacts", reply_markup=keyboard)
                await message.answer("Добро пожаловать в нашего помощника по заполнению технических заданий!\n\n📝 Заполнить ТЗ — Нажмите /technical_specification, чтобы начать или продолжить\n\n🤖 Хелпер — Получите помощь и советы по использованию бота /help", reply_markup=keyboard)
            else:
                await message.answer("Владелец бота не продлил свою подписку, функционал временно ограничен.\nБот разработан @WSCHUDO")
    

    
    # Start Модель Портфолио

    @dp.message_handler(commands=['portfolio'])
    @dp.message_handler(lambda message: message.text == "Портфолио")
    async def portfolio_start(message: types.Message, state: FSMContext):
        subscription_status = await check_subscription(message.from_user.id)

        if subscription_status['status']:

            await state.update_data(page=1)
            await show_portfolio_page(message, 1)

        else:
            await message.answer(subscription_status['message'])      

    async def show_portfolio_page(message: types.Message, page: int, edit: bool = False):
        

        if page == 1:
            text = "Первая страница портфолио"
            text += "\n[Изображение](https://example.com/image.jpg)"
            text += "\n[Видео](https://example.com/video.mp4)"
            keyboard = types.InlineKeyboardMarkup()
            keyboard.add(types.InlineKeyboardButton("Далее ➡️", callback_data="portfolio_next"))
            keyboard.add(types.InlineKeyboardButton("❌ Закрыть ❌", callback_data="portfolio_close"))
        elif page == 2:
            text = "Вторая страница портфолио"
            text += "\n[Изображение](https://example.com/image.jpg)"
            text += "\n[Видео](https://example.com/video.mp4)"
            keyboard = types.InlineKeyboardMarkup()
            keyboard.row(
                types.InlineKeyboardButton("⬅️ Назад", callback_data="portfolio_prev"),
                types.InlineKeyboardButton("Далее ➡️", callback_data="portfolio_next")
            )
            keyboard.add(types.InlineKeyboardButton("❌ Закрыть ❌", callback_data="portfolio_close"))
        elif page == 3:
            text = "Третья страница портфолио"
            text += "\n[Изображение](https://example.com/image.jpg)"
            text += "\n[Видео](https://example.com/video.mp4)"
            keyboard = types.InlineKeyboardMarkup()
            keyboard.add(types.InlineKeyboardButton("⬅️ Назад", callback_data="portfolio_prev"))
            keyboard.add(types.InlineKeyboardButton("❌ Закрыть ❌", callback_data="portfolio_close"))

        if edit:
            await message.edit_text(text, reply_markup=keyboard, parse_mode=types.ParseMode.MARKDOWN)
        else:
            await message.answer(text, reply_markup=keyboard, parse_mode=types.ParseMode.MARKDOWN)

                

    @dp.callback_query_handler(lambda c: c.data in ['portfolio_next', 'portfolio_prev', 'portfolio_close'])
    async def handle_portfolio_callback(call: types.CallbackQuery, state: FSMContext):
        user_data = await state.get_data()
        current_page = user_data.get('page', 1)
        
        if call.data == 'portfolio_next':
            next_page = current_page + 1
            await state.update_data(page=next_page)
            await show_portfolio_page(call.message, next_page, edit=True)
        elif call.data == 'portfolio_prev':
            prev_page = current_page - 1
            await state.update_data(page=prev_page)
            await show_portfolio_page(call.message, prev_page, edit=True)
        elif call.data == 'portfolio_close':
            await call.message.delete()
            await state.finish()
        await call.answer()


    # End Модель Портфолио

    @dp.message_handler(commands=['help'])
    @dp.message_handler(lambda message: message.text == "Хелпер")
    async def helper_command(message: types.Message):

        helper_text = (
            "Добро пожаловать в раздел 'Хелпер'!\n\n"
            "Здесь вы можете найти ответы на часто задаваемые вопросы и получить помощь по использованию бота.\n\n"
            "1. Как использовать бота?\n"
            "2. Как заполнить техническое задание?\n"
            "4. Как связаться с поддержкой?\n\n"
            "Если у вас есть другие вопросы, пожалуйста, свяжитесь с нашей поддержкой."
        )
        await message.answer(helper_text)
 
    @dp.message_handler(commands=['contacts'])
    @dp.message_handler(lambda message: message.text == "Контакты")
    async def contacts_command(message: types.Message):
        subscription_status = await check_subscription(message.from_user.id)

        if subscription_status['status']:

            contacts_text = (
                "Контакты:\n\n"
                "1. Email: support@example.com\n"
                "2. Телефон: +1 234 567 8900\n"
                "3. Адрес: 123 Main Street, Anytown, AT 12345\n\n"
                "Если у вас есть вопросы или вам нужна помощь, пожалуйста, свяжитесь с нами по указанным контактам."
            )
            await message.answer(contacts_text)

        else:
            await message.answer(subscription_status['message'])  

    # Заполнение ТЗ --------------------------------
    
    @dp.message_handler(commands=['technical_specification'])
    @dp.message_handler(lambda message: message.text == "Заполнить ТЗ", state='*')
    async def ask_for_phone(message: types.Message, state: FSMContext):
        async with db_pool.acquire() as connection:
            # Проверяем статус пользователя в таблице users_designer
            user_status = await connection.fetchval(
                """
                SELECT status FROM users_designer
                WHERE id_telegram = $1
                """,
                message.from_user.id
            )

            if user_status == 1:
                # Пользователь уже начал опрос, но не завершил
                last_step = await connection.fetchval(
                    """
                    SELECT last_step FROM users_designer
                    WHERE id_telegram = $1
                    """,
                    message.from_user.id
                )
                request_id = await get_request_id(message.from_user.id)
                # Переходим к вопросу, который был последним
                await start_questionnaire(message, state, last_step, request_id)
            else:
                # Запрашиваем номер телефона, если опрос еще не начинался
                keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True)
                button_phone = types.KeyboardButton(text="Поделиться номером телефона", request_contact=True)
                button_cancel = types.KeyboardButton(text="Нет")
                keyboard.add(button_phone, button_cancel)
                await Form.waiting_for_phone.set()
                await message.answer("Пожалуйста, поделитесь своим номером телефона или нажмите 'Нет', если не хотите продолжать.", reply_markup=keyboard)


    @dp.message_handler(content_types=types.ContentTypes.CONTACT, state=Form.waiting_for_phone)
    async def phone_received(message: types.Message, state: FSMContext):
        user_data = await state.get_data()
        phone_number = message.contact.phone_number
        user_data['phone'] = phone_number  # Сохраняем телефон во временное хранилище состояний
        
        async with db_pool.acquire() as connection:
            # Начинаем транзакцию
            async with connection.transaction():
                # Попытка получить данные пользователя из таблицы users_designer по ID телеграма
                user_info = await connection.fetchrow(
                    "SELECT * FROM users_designer WHERE id_telegram = $1",
                    message.from_user.id
                )
                
                if user_info:
                    # Обновляем данные в users_designer
                    await connection.execute(
                        """
                        UPDATE users_designer SET
                            phone = $1,
                            status = 1
                        WHERE id_telegram = $2
                        """,
                        phone_number,
                        message.from_user.id
                    )

                    # Записываем данные в таблицу data_questions
                    await connection.execute(
                        """
                        INSERT INTO data_questions (
                            user_id, id_telegram, tg_login, tg_firstname, tg_lastname, phone, root
                        ) VALUES ($1, $2, $3, $4, $5, $6, $7)
                        """,
                        user_info['id_telegram'],  # Или другой уникальный ID, если это необходимо
                        message.from_user.id,
                        message.from_user.username,
                        message.from_user.first_name,
                        message.from_user.last_name,
                        phone_number,
                        get_folder_id()  # Это функция, которая должна быть определена ранее для извлечения ID из названия папки
                    )
                    request_id = await get_request_id(message.from_user.id)
                    await state.update_data(db_record_created=True)  # Флаг, что запись создана
                    keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True)
                    buttons = ["Продолжить", "Позже"]
                    keyboard.row(*buttons)
                    await Form.waiting_for_phone_confirmation.set()
                    await message.answer("Ваш номер сохранен. Хотите продолжить?", reply_markup=keyboard)
                else:
                    await message.answer("Не удалось найти вашу учетную запись в системе.", reply_markup=types.ReplyKeyboardRemove())
                    await state.finish()  # Завершаем состояние, если информация о пользователе не найдена



    @dp.message_handler(lambda message: message.text == "Нет", state=Form.waiting_for_phone)
    async def operation_cancelled(message: types.Message, state: FSMContext):
        await message.answer("Операция прервана.", reply_markup=types.ReplyKeyboardRemove())
        await state.finish()

    @dp.message_handler(lambda message: message.text == "Позже", state=Form.waiting_for_phone_confirmation)
    async def operation_cancelled(message: types.Message, state: FSMContext):
        await message.answer("Операция прервана.", reply_markup=types.ReplyKeyboardRemove())
        await state.finish()

    # Продолжение    

    async def load_user_answers(user_id, request_id):
        async with db_pool.acquire() as connection:
            rows = await connection.fetch("""
                SELECT question_step, answer_text, answer_type FROM user_answers
                WHERE id_telegram = $1 AND request_id = $2
            """, user_id, request_id)
            answers = {}
            custom_answers = {}
            for row in rows:
                step = row['question_step']
                answer = row['answer_text']
                if row['answer_type'] == 'custom':
                    custom_answers[step] = answer
                else:
                    if step not in answers:
                        answers[step] = []
                    answers[step].append(answer)
            return answers, custom_answers

    @dp.message_handler(lambda message: message.text == "Продолжить", state='*')
    async def start_questionnaire(message: types.Message, state: FSMContext, last_step=None, request_id=None):

        subscription_status = await check_subscription(message.from_user.id)

        if subscription_status['status']:

            questions = load_questions()
            user_id = message.from_user.id

            if request_id is None:
                request_id = await get_request_id(user_id)

            # Загружаем сохраненные ответы
            answers, custom_answers = await load_user_answers(user_id, request_id)

            if last_step is None:
                # Получаем последний сохраненный шаг из базы данных
                async with db_pool.acquire() as connection:
                    last_step = await connection.fetchval(
                        """
                        SELECT last_step FROM users_designer
                        WHERE id_telegram = $1
                        """,
                        user_id
                    )

            # Устанавливаем индекс текущего вопроса, начиная с последнего сохраненного шага
            current_question_index = max(last_step - 1, 0)

            # Обновляем состояние
            await state.update_data(questions=questions, current_question_index=current_question_index, answers=answers, custom_answers=custom_answers, request_id=request_id)
            
            await ask_question(message, state)
            
        else:
            await message.answer(subscription_status['message'])  

    # Функция для обрезки текста
    def truncate_text(text, max_length=64):
        encoded_text = text.encode('utf-8')
        if len(encoded_text) > max_length:
            truncated_text = encoded_text[:max_length].decode('utf-8', errors='ignore')
            if len(truncated_text) < len(text):
                truncated_text = truncated_text[:max_length//2] + "..." + truncated_text[max_length//2+3:]
            return truncated_text
        return text

    # Функция для создания клавиатуры
    async def create_keyboard(current_index, selected_answers, questions, custom_answers):
        keyboard = types.InlineKeyboardMarkup(row_width=1)
        options = questions[current_index].get("options", [])
        for option in options:
            option_text = option["text"]
            display_text = f"{option_text} ✅" if option_text in selected_answers else option_text
            # Обрезаем текст, только чтобы учесть длину 'answer:{current_index}:'
            truncated_text = truncate_text(option['text'], 62 - len(f'answer:{current_index}:'))
            callback_data = f"answer:{current_index}:{truncated_text}"
            keyboard.add(types.InlineKeyboardButton(text=display_text, callback_data=callback_data))

        custom_text = "Свой вариант ✅" if custom_answers.get(current_index) else "Свой вариант"
        keyboard.add(types.InlineKeyboardButton(text=custom_text, callback_data="custom_answer"))

        nav_buttons = []
        if current_index > 0:
            nav_buttons.append(types.InlineKeyboardButton(text="⬅️ Назад", callback_data="nav:back"))
        if current_index < len(questions) - 1:
            nav_buttons.append(types.InlineKeyboardButton(text="Вперёд ➡️", callback_data="nav:forward"))

        if nav_buttons:
            keyboard.row(*nav_buttons)

        if current_index == len(questions) - 1:
            keyboard.add(types.InlineKeyboardButton(text="❎ Завершить опрос ❎", callback_data="nav:end"))
            
        keyboard.add(types.InlineKeyboardButton(text="❌ Прерваться ❌", callback_data="nav:interrupt"))

        return keyboard


    # Упрощённый код для ask_question
    async def ask_question(message: types.Message, state: FSMContext):
        user_data = await state.get_data()
        current_index = user_data['current_question_index']
        question_info = user_data['questions'][current_index]
        custom_answers = user_data.get('custom_answers', {})
        answers = user_data.get('answers', {})
        selected_answers = answers.get(current_index, [])

        # Удаляем предыдущую клавиатуру
        await message.answer(question_info["text"], reply_markup=types.ReplyKeyboardRemove())


        keyboard = await create_keyboard(current_index, selected_answers, user_data['questions'], custom_answers)

        # Проверка наличия изображений и их отправка в виде галереи
        if 'options' in question_info and any('image' in option for option in question_info['options']):
            caption_text = f"Изображения к вопросу {current_index + 1}: {question_info['text']}"
            media_group = [types.InputMediaPhoto(option['image'], caption=caption_text if i == 0 else None) for i, option in enumerate(question_info['options']) if 'image' in option]
            await message.answer_media_group(media_group)

        await Questionnaire.asking.set()
        await message.answer(question_info["text"], reply_markup=keyboard)

    async def handle_answer(callback_query: types.CallbackQuery, state: FSMContext):
        data = callback_query.data.split(':')
        action = data[0]
        user_data = await state.get_data()
        questions = user_data['questions']
        current_index = user_data['current_question_index']
        request_id = user_data['request_id']
        question_key = questions[current_index]["key"]
        answers = user_data.get('answers', {})
        selected_answers = answers.get(current_index, [])
        user_id = callback_query.from_user.id
        tg_login = callback_query.from_user.username

        if action == "answer":
            current_index, truncated_answer = int(data[1]), data[2]
            # Найдём полный текст ответа по обрезанному тексту
            full_answer = next((opt['text'] for opt in questions[current_index]["options"] if truncate_text(opt['text'], 62 - len(f'answer:{current_index}:')) == truncated_answer), truncated_answer)
            response_message = "Ответ сохранён"  # Сообщение по умолчанию

            answer_type = "button" if full_answer in [opt['text'] for opt in questions[current_index].get("options", [])] else "custom"
            if full_answer in selected_answers:
                selected_answers.remove(full_answer)
                await remove_user_answer_from_db(user_id, request_id, current_index, full_answer)
                response_message = "Ответ удалён"
            else:
                selected_answers.append(full_answer)
                await save_user_answer_to_db(user_id, tg_login, request_id, current_index, full_answer, answer_type)
                response_message = "Ответ сохранён"

            answers[current_index] = selected_answers
            user_data['answers'] = answers
            await state.update_data(user_data)
            await update_question_message(callback_query.message, current_index, state, questions, update_images=False)
            await update_step_in_database(callback_query.from_user.id, current_index)  # Обновляем шаг пользователя в базе
            await callback_query.answer(response_message)  # Отправляем корректное сообщение пользователю

        elif action == "nav":
            direction = data[1]
            update_images = False

            if direction == "forward" and current_index < len(questions) - 1:
                current_index += 1
                update_images = True
            elif direction == "back" and current_index > 0:
                current_index -= 1
                update_images = True
            elif direction == "end":
                await finish_questionnaire(callback_query, state)
                return
            elif direction == "interrupt":
                await interrupt_questionnaire(callback_query, state)
                return

            user_data['current_question_index'] = current_index
            await state.update_data(user_data)
            await update_question_message(callback_query.message, current_index, state, questions, update_images=update_images)
            await update_step_in_database(callback_query.from_user.id, current_index)  # Обновляем шаг пользователя в базе
            await callback_query.answer()



    @dp.message_handler(state=Questionnaire.custom_answer)
    async def process_custom_answer(message: types.Message, state: FSMContext):
        user_data = await state.get_data()
        current_index = user_data['current_question_index']
        request_id = user_data['request_id']
        questions = user_data['questions']
        user_id = message.from_user.id
        tg_login = message.from_user.username
        question_key = questions[current_index]["key"]
        custom_answers = user_data.get('custom_answers', {})
        answer_text = message.text
        answer_type = "custom"

        # Сохраняем или обновляем ответ в словаре
        custom_answers[current_index] = answer_text
        await state.update_data(custom_answers=custom_answers)

        # Сохраняем ответ в базу данных
        await save_custom_answer_to_db(user_id, tg_login, request_id, current_index, question_key, answer_text, answer_type)

        # Удаление сообщения пользователя и переотправка вопроса
        await message.delete()
        await ask_question(message, state)
        await Questionnaire.asking.set()

        

    @dp.callback_query_handler(text="custom_answer", state=Questionnaire.asking)
    async def handle_custom_answer(callback_query: types.CallbackQuery, state: FSMContext):
        user_data = await state.get_data()
        current_index = user_data['current_question_index']
        request_id = user_data['request_id']
        questions = user_data['questions']
        user_id = callback_query.from_user.id
        question_key = questions[current_index]["key"]
        custom_answers = user_data.get('custom_answers', {})

        if custom_answers.get(current_index):
            # Удалить пользовательский ответ из базы данных
            await remove_custom_answer_from_db(user_id, request_id, current_index, custom_answers[current_index])

            # Удаляем ответ из словаря
            del custom_answers[current_index]
            await state.update_data(custom_answers=custom_answers)

            await update_question_message(callback_query.message, current_index, state, questions, update_images=False)
            await callback_query.answer("Ответ удалён")
        else:
            # Переход в состояние ввода пользовательского ответа
            await Questionnaire.custom_answer.set()
            await callback_query.message.delete()
            await callback_query.message.answer("Введите ваш ответ:")





    # Использование общей функции для создания клавиатуры в update_question_message
    async def update_question_message(message: types.Message, current_index: int, state: FSMContext, questions, update_images=False):
        user_data = await state.get_data()
        question_info = questions[current_index]
        custom_answers = user_data.get('custom_answers', {})
        answers = user_data.get('answers', {})
        selected_answers = answers.get(current_index, [])

        keyboard = await create_keyboard(current_index, selected_answers, questions, custom_answers)
        new_text = question_info["text"]

        try:
            if update_images and 'options' in question_info and any('image' in option for option in question_info['options']):
                caption_text = f"Изображения к вопросу {current_index + 1}: {new_text}"
                media_group = [
                    types.InputMediaPhoto(option['image'], caption=caption_text if i == 0 else None)
                    for i, option in enumerate(question_info['options']) if 'image' in option
                ]
                await message.delete()
                await message.answer_media_group(media_group)
                await message.answer(new_text, reply_markup=keyboard)
            else:
                # Сравниваем текущее состояние сообщения с новым
                if message.text != new_text or message.reply_markup != keyboard:
                    await message.edit_text(text=new_text, reply_markup=keyboard)
        except MessageNotModified:
            pass







    async def update_step_in_database(user_id, current_index):
        async with db_pool.acquire() as connection:
            await connection.execute(
                """
                UPDATE users_designer SET last_step = $1 WHERE id_telegram = $2
                """,
                current_index + 1, user_id
            )
            await connection.execute(
                """
                UPDATE data_questions SET step_number = $1, step_time = CURRENT_TIMESTAMP
                WHERE id_telegram = $2 AND step_time = (
                    SELECT MAX(step_time) FROM data_questions WHERE id_telegram = $2
                )
                """,
                current_index + 1, user_id
            )

    # Завершение опроса и обновление статуса в базе данных
    async def finish_questionnaire(callback_query: types.CallbackQuery, state: FSMContext):
        user_id = callback_query.from_user.id
        await callback_query.message.answer("Спасибо за ваши ответы!", reply_markup=types.ReplyKeyboardRemove())
        
        # Удаляем inline-клавиатуру из последнего сообщения с вопросом
        if callback_query.message:
            await callback_query.message.edit_reply_markup(reply_markup=None)
        
        async with db_pool.acquire() as connection:
            await connection.execute(
                """
                UPDATE users_designer SET status = 0, last_step = 0 WHERE id_telegram = $1
                """,
                user_id
            )
            await connection.execute(
                """
                UPDATE data_questions SET step_number = -1, step_time = CURRENT_TIMESTAMP
                WHERE id_telegram = $1 AND step_time = (
                    SELECT MAX(step_time) FROM data_questions WHERE id_telegram = $1
                )
                """,
                user_id
            )
        
        await state.finish()
        
        # Создаем и отправляем документ администратору после завершения опроса
        request_id = await get_request_id(user_id)
        file_path = await create_word_document(user_id, request_id)
        admin_id = get_folder_id()
        
        admin_message = f"Уважаемый администратор, поступила новая заявка от пользователя @{callback_query.from_user.username} {callback_query.from_user.first_name} {callback_query.from_user.last_name}"
        if not callback_query.from_user.username:
            admin_message = f"Уважаемый администратор, поступила новая заявка от пользователя {callback_query.from_user.first_name} {callback_query.from_user.last_name}"
        
        await bot.send_message(admin_id, admin_message)
        await bot.send_document(admin_id, open(file_path, 'rb'))





    # Сохранение ответов в БД
    
    async def save_user_answer_to_db(user_id, tg_login, request_id, question_step, answer_text, answer_type):
        root_id = await get_root_id(user_id)

        async with db_pool.acquire() as connection:
            await connection.execute("""
                INSERT INTO user_answers (
                    id_telegram, tg_login, request_id, question_step, answer_text, answer_type, root
                ) VALUES ($1, $2, $3, $4, $5, $6, $7)
            """, user_id, tg_login, request_id, question_step, answer_text, answer_type, root_id)

    async def save_custom_answer_to_db(user_id, tg_login, request_id, question_step, question_key, answer_text, answer_type):
        root_id = await get_root_id(user_id)

        async with db_pool.acquire() as connection:
            await connection.execute("""
                INSERT INTO user_answers (
                    id_telegram, tg_login, request_id, question_step, answer_text, answer_type, root
                ) VALUES ($1, $2, $3, $4, $5, $6, $7)
            """, user_id, tg_login, request_id, question_step, answer_text, answer_type, root_id)

    async def remove_custom_answer_from_db(user_id, request_id, question_step, answer_text):
        async with db_pool.acquire() as connection:
            await connection.execute("""
                DELETE FROM user_answers
                WHERE id_telegram = $1 AND request_id = $2 AND question_step = $3 AND answer_text = $4
            """, user_id, request_id, question_step, answer_text)

    async def remove_user_answer_from_db(user_id, request_id, question_step, answer_text):
        async with db_pool.acquire() as connection:
            await connection.execute("""
                DELETE FROM user_answers WHERE id_telegram=$1 AND request_id=$2 AND question_step=$3 AND answer_text=$4
            """, user_id, request_id, question_step, answer_text)

    async def get_root_id(user_id):
        async with db_pool.acquire() as connection:
            root_id = await connection.fetchval("""
                SELECT root FROM users_designer WHERE id_telegram = $1
            """, user_id)
            return root_id

    async def get_request_id(user_id):
        async with db_pool.acquire() as connection:
            # Получаем ID последней заявки пользователя, ориентируясь по дате начала заявки (step_start).
            request_id = await connection.fetchval("""
                SELECT id FROM data_questions 
                WHERE id_telegram = $1
                ORDER BY step_start DESC
                LIMIT 1
            """, user_id)
            return request_id

    # Формирование doсx

    # Создание документа Word
    async def create_word_document(user_id, request_id):
        async with db_pool.acquire() as connection:
            # Получаем данные пользователя и опроса
            data_question = await connection.fetchrow("""
                SELECT * FROM data_questions
                WHERE id_telegram = $1 AND id = $2
            """, user_id, request_id)
            
            # Получаем ответы пользователя
            user_answers = await connection.fetch("""
                SELECT * FROM user_answers
                WHERE id_telegram = $1 AND request_id = $2
            """, user_id, request_id)
            
        # Создаем документ
        doc = Document()
        doc.add_heading('Отчет по опросу', 0)

        # Форматируем даты
        step_start = datetime.strptime(str(data_question['step_start']), '%Y-%m-%d %H:%M:%S.%f%z').strftime('%d.%m.%Y %H:%M')
        step_time = datetime.strptime(str(data_question['step_time']), '%Y-%m-%d %H:%M:%S.%f%z').strftime('%d.%m.%Y %H:%M')

        # Добавляем данные пользователя и опроса
        doc.add_paragraph(f"ID заявки: {request_id}")
        doc.add_paragraph(f"Дата начала прохождения опроса: {step_start}")
        doc.add_paragraph(f"Дата последнего ответа: {step_time}")
        doc.add_paragraph(f"Логин пользователя: {data_question['tg_login']}")
        doc.add_paragraph(f"Фамилия и Имя пользователя: {data_question['tg_firstname']} {data_question['tg_lastname']}")
        doc.add_paragraph(f"Телефон пользователя: {data_question['phone']}")

        # Добавляем ответы пользователя
        doc.add_heading('Ответы пользователя', level=1)
        
        # Получаем вопросы
        questions = load_questions()
        questions_dict = {i: q for i, q in enumerate(questions)}

        # Преобразуем ответы в словарь для быстрого доступа
        answers_dict = {(answer['question_step'], answer['answer_text']): answer for answer in user_answers}

        for step, question_info in questions_dict.items():
            question_text = question_info['text']
            doc.add_heading(f"Вопрос {step + 1}: {question_text}", level=2)
            
            # Проверяем ответы на текущий вопрос
            question_options = question_info.get('options', [])
            has_standard_answer = False

            for option in question_options:
                answer_text = option['text']
                answer_key = (step, answer_text)
                if answer_key in answers_dict:
                    answer = answers_dict[answer_key]
                    answer_type = '(стандартный)' if answer['answer_type'] == 'button' else '(пользовательский)'
                    doc.add_paragraph(f"Ответ: {answer_text} {answer_type}")
                    has_standard_answer = True

                    # Включаем изображение, если это ответ с изображением
                    if 'image' in option:
                        response = requests.get(option['image'])
                        image_stream = BytesIO(response.content)
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run()
                        run.add_picture(image_stream, width=Inches(1), height=Inches(1))
            
            # Проверяем пользовательский ответ на текущий вопрос
            custom_answer = next((a for a in user_answers if a['question_step'] == step and a['answer_type'] == 'custom'), None)
            if custom_answer:
                doc.add_paragraph(f"Ответ: {custom_answer['answer_text']} (пользовательский)")

            # Если ответов на текущий вопрос нет
            if not has_standard_answer and not custom_answer:
                doc.add_paragraph("Пользователь не ответил на этот вопрос.")

        # Определяем путь сохранения документа
        folder_name = get_folder_id()
        folder_path = f"user_bots/{folder_name}/data_questions"
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)

        file_name = f"{data_question['tg_login'] or data_question['tg_firstname'] + ' ' + data_question['tg_lastname']} {user_id} {request_id}.docx"
        file_path = os.path.join(folder_path, file_name)
        
        doc.save(file_path)
        
        return file_path

    # Прерывание процесса

    # Функция для прерывания опроса
    async def interrupt_questionnaire(callback_query: types.CallbackQuery, state: FSMContext):
        user_id = callback_query.from_user.id
        
        # Удаляем сообщение с вариантами ответа
        await callback_query.message.delete()
        
        # Сбрасываем состояние пользователя
        await state.finish()
        
        # Создаем клавиатуру с кнопками "Продолжить" и "Меню"
        keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True)
        keyboard.add(types.KeyboardButton(text="Продолжить"))
        keyboard.add(types.KeyboardButton(text="Меню"))
        
        # Отправляем пользователю сообщение о возможности продолжения заполнения ТЗ
        await callback_query.message.answer("Вы всегда можете продолжить заполнение технического задания", reply_markup=keyboard)


    # END Прерывание процесса

    # Обработчик команды /admin
    # @dp.message_handler(commands=['admin'])
    # async def admin_menu(message: types.Message):
    #     admin_id = await get_folder_id()  # Если get_folder_id асинхронная, используем await
        
    #     if message.from_user.id == admin_id:
    #         # Создание клавиатуры для администратора
    #         keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
    #         buttons = [
    #             ["Мои файлы", "Мои пользователи"],
    #             ["Портфолио", "Контакты"]
    #         ]
    #         for row in buttons:
    #             keyboard.row(*row)

    #         await message.answer("Административное меню:", reply_markup=keyboard)
    #     else:
    #         await message.answer("У вас нет доступа к этому разделу.")
    

    dp.register_callback_query_handler(handle_answer, Text(startswith=["answer:", "nav:"]), state=Questionnaire.asking)

    

    # Запуск бота
    await dp.start_polling()

if __name__ == '__main__':
    asyncio.run(main())
